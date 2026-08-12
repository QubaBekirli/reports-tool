from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import os
from xml.sax.saxutils import escape as _xml_escape


def _esc(text: str) -> str:
    return _xml_escape(text or "")


_FONT_REGISTERED = False

def _register_fonts():
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return
    dejavu = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    dejavu_bold = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if os.path.exists(dejavu):
        pdfmetrics.registerFont(TTFont("DejaVu", dejavu))
        pdfmetrics.registerFont(TTFont("DejaVu-Bold", dejavu_bold))
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        registerFontFamily("DejaVu", normal="DejaVu", bold="DejaVu-Bold",
                            italic="DejaVu", boldItalic="DejaVu-Bold")
        _FONT_REGISTERED = True


STATUS_COLORS_DOCX = {
    "compliant": RGBColor(0x10, 0x99, 0x68),
    "partial": RGBColor(0xF5, 0x9E, 0x0B),
    "missing": RGBColor(0xEF, 0x44, 0x44),
}

STATUS_COLORS_PDF = {
    "compliant": HexColor("#10B981"),
    "partial": HexColor("#F59E0B"),
    "missing": HexColor("#EF4444"),
}

STATUS_LABELS = {
    "az": {"compliant": "Uyğundur", "partial": "Qismən uyğundur", "missing": "Yoxdur"},
    "en": {"compliant": "Compliant", "partial": "Partially Compliant", "missing": "Missing"},
}

RISK_LABELS = {
    "az": {"low": "Aşağı", "medium": "Orta", "high": "Yüksək", "critical": "Kritik"},
    "en": {"low": "Low", "medium": "Medium", "high": "High", "critical": "Critical"},
}


def _get_lang(result: dict) -> str:
    lang = result.get("detected_language", "az")
    return lang if lang in ("az", "en") else "az"


def generate_docx(result: dict) -> bytes:
    lang = _get_lang(result)
    labels = STATUS_LABELS[lang]
    risk_labels = RISK_LABELS[lang]
    classification = result.get("document_classification") or {}

    doc = DocxDocument()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Gap-Analiz Hesabatı" if lang == "az" else "Gap Analysis Report", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(result.get("document_title", ""))
    doc.add_paragraph(
        (f"Risk səviyyəsi: {risk_labels.get(result.get('risk_level', 'medium'), 'Orta')}")
        if lang == "az"
        else f"Risk level: {risk_labels.get(result.get('risk_level', 'medium'), 'Medium')}"
    )

    # Document Classification
    if classification and classification.get("detected_type"):
        doc.add_heading(
            "Sənəd Klassifikasiyası" if lang == "az" else "Document Classification", level=1
        )
        if lang == "az":
            doc.add_paragraph(f"Növ: {classification.get('detected_type', '')}")
            doc.add_paragraph(f"Tətbiq olunan standart: {classification.get('applicable_standard', '')}")
            clauses = classification.get("primary_clauses", [])
            if clauses:
                doc.add_paragraph(f"Əsas bendlər: {', '.join(clauses)}")
        else:
            doc.add_paragraph(f"Type: {classification.get('detected_type', '')}")
            doc.add_paragraph(f"Applicable standard: {classification.get('applicable_standard', '')}")
            clauses = classification.get("primary_clauses", [])
            if clauses:
                doc.add_paragraph(f"Primary clauses: {', '.join(clauses)}")

    doc.add_heading(
        "İcraçı Xülasə" if lang == "az" else "Executive Summary", level=1
    )
    doc.add_paragraph(result.get("executive_summary", ""))

    doc.add_heading(
        "Statistika" if lang == "az" else "Statistics", level=1
    )
    stats_table = doc.add_table(rows=2, cols=4)
    stats_table.style = "Light Grid Accent 1"
    headers = ["Uyğun", "Qismən", "Yoxdur", "Cəmi"] if lang == "az" else ["Compliant", "Partial", "Missing", "Total"]
    for i, h in enumerate(headers):
        stats_table.rows[0].cells[i].text = h
    stats_table.rows[1].cells[0].text = str(result.get("compliant_count", 0))
    stats_table.rows[1].cells[1].text = str(result.get("partial_count", 0))
    stats_table.rows[1].cells[2].text = str(result.get("missing_count", 0))
    stats_table.rows[1].cells[3].text = str(result.get("total_controls", 0))

    # Gap details grouped by source
    gaps_by_source: dict[str, list] = {}
    for gap in result.get("gaps", []):
        src = gap.get("control_source", "")
        gaps_by_source.setdefault(src, []).append(gap)

    doc.add_heading(
        "Gap Detalları" if lang == "az" else "Gap Details", level=1
    )
    for source, gaps in gaps_by_source.items():
        doc.add_heading(source, level=2)
        for gap in gaps:
            status = gap.get("status", "missing")
            status_label = labels.get(status, status)
            p = doc.add_paragraph()
            run = p.add_run(
                f"[{status_label}] {gap.get('control_category', '')} "
                f"({gap.get('control_id', '')})"
            )
            run.bold = True
            run.font.color.rgb = STATUS_COLORS_DOCX.get(status, RGBColor(0, 0, 0))

            if gap.get("control_text"):
                ctrl_label = "Nəzarət tələbi" if lang == "az" else "Control requirement"
                doc.add_paragraph(f"{ctrl_label}: {gap['control_text']}")

            if gap.get("standard_requirement"):
                req_label = "Standart tələbi" if lang == "az" else "Standard requirement"
                doc.add_paragraph(f"{req_label}: {gap['standard_requirement']}")

            cur_text = gap.get("current_document_text", "")
            if cur_text and cur_text != "Text not found":
                cur_label = "Sənəddəki mətn" if lang == "az" else "Current document text"
                doc.add_paragraph(f"{cur_label}: {cur_text}")

            if gap.get("gap_analysis"):
                gap_label = "Gap analizi" if lang == "az" else "Gap analysis"
                doc.add_paragraph(f"{gap_label}: {gap['gap_analysis']}")

            if gap.get("remediation_proposal"):
                rem_label = "Remediation təklifi" if lang == "az" else "Remediation proposal"
                doc.add_paragraph(f"{rem_label}: {gap['remediation_proposal']}")

            if gap.get("potential_risks"):
                risk_label = "Potensial risklər" if lang == "az" else "Potential risks"
                doc.add_paragraph(f"{risk_label}: {gap['potential_risks']}")

            if gap.get("justification"):
                just_label = "Əsaslandırma" if lang == "az" else "Justification"
                doc.add_paragraph(f"{just_label}: {gap.get('justification', '')}")

            if gap.get("evidence_snippet"):
                ev_label = "Dəlil" if lang == "az" else "Evidence"
                ev_ref = gap.get("evidence_reference", "")
                ev_text = f"{ev_label}: {gap['evidence_snippet']}"
                if ev_ref:
                    ev_text += f" ({ev_ref})"
                doc.add_paragraph(ev_text, style="Intense Quote")

    doc.add_heading(
        "Tövsiyələr" if lang == "az" else "Recommendations", level=1
    )
    for i, rec in enumerate(result.get("recommendations", []), 1):
        doc.add_paragraph(f"{i}. {rec}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_pdf(result: dict) -> bytes:
    lang = _get_lang(result)
    labels = STATUS_LABELS[lang]
    risk_labels = RISK_LABELS[lang]
    classification = result.get("document_classification") or {}

    _register_fonts()
    font_name = "DejaVu" if _FONT_REGISTERED else "Helvetica"
    font_bold = "DejaVu-Bold" if _FONT_REGISTERED else "Helvetica-Bold"

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, margins=[2 * cm, 2 * cm, 2 * cm, 2 * cm])
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle", parent=styles["Title"], fontName=font_bold,
        fontSize=20, spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "CustomHeading", parent=styles["Heading2"], fontName=font_bold,
        fontSize=14, spaceBefore=12, spaceAfter=6,
    )
    sub_heading_style = ParagraphStyle(
        "CustomSubHeading", parent=styles["Heading3"], fontName=font_bold,
        fontSize=11, spaceBefore=8, spaceAfter=4, textColor=HexColor("#444444"),
    )
    body_style = ParagraphStyle(
        "CustomBody", parent=styles["Normal"], fontName=font_name,
        fontSize=10, leading=14, spaceAfter=6,
    )
    evidence_style = ParagraphStyle(
        "Evidence", parent=body_style, fontName=font_name,
        leftIndent=20, textColor=HexColor("#666666"), spaceAfter=6,
    )
    remediation_style = ParagraphStyle(
        "Remediation", parent=body_style, fontName=font_name,
        leftIndent=12, textColor=HexColor("#7C3AED"), spaceAfter=6,
    )

    elements = []

    elements.append(Paragraph(
        "Gap-Analiz Hesabatı" if lang == "az" else "Gap Analysis Report", title_style
    ))
    elements.append(Paragraph(_esc(result.get("document_title", "")), body_style))
    elements.append(Paragraph(
        (f"Risk səviyyəsi: {risk_labels.get(result.get('risk_level', 'medium'), 'Orta')}")
        if lang == "az"
        else f"Risk level: {risk_labels.get(result.get('risk_level', 'medium'), 'Medium')}",
        body_style,
    ))
    elements.append(Spacer(1, 12))

    # Document Classification
    if classification and classification.get("detected_type"):
        elements.append(Paragraph(
            "Sənəd Klassifikasiyası" if lang == "az" else "Document Classification",
            heading_style,
        ))
        type_label = "Aşkar edilmiş növ" if lang == "az" else "Detected type"
        std_label = "Tətbiq olunan standart" if lang == "az" else "Applicable standard"
        clauses_label = "Əsas bendlər" if lang == "az" else "Primary clauses"
        elements.append(Paragraph(
            f"<b>{type_label}:</b> {_esc(classification.get('detected_type', ''))}", body_style
        ))
        elements.append(Paragraph(
            f"<b>{std_label}:</b> {_esc(classification.get('applicable_standard', ''))}", body_style
        ))
        clauses = classification.get("primary_clauses", [])
        if clauses:
            elements.append(Paragraph(
                f"<b>{clauses_label}:</b> {_esc(', '.join(clauses))}", body_style
            ))
        elements.append(Spacer(1, 8))

    # Executive Summary
    elements.append(Paragraph(
        "İcraçı Xülasə" if lang == "az" else "Executive Summary", heading_style
    ))
    elements.append(Paragraph(
        _esc(result.get("executive_summary", "")).replace("\n", "<br/>"), body_style
    ))
    elements.append(Spacer(1, 8))

    # Statistics
    elements.append(Paragraph(
        "Statistika" if lang == "az" else "Statistics", heading_style
    ))
    stats_headers = (
        ["Uyğun", "Qismən", "Yoxdur", "Cəmi"]
        if lang == "az" else ["Compliant", "Partial", "Missing", "Total"]
    )
    stats_data = [
        stats_headers,
        [
            str(result.get("compliant_count", 0)),
            str(result.get("partial_count", 0)),
            str(result.get("missing_count", 0)),
            str(result.get("total_controls", 0)),
        ],
    ]
    stats_table = Table(stats_data, colWidths=[3 * cm, 3 * cm, 3 * cm, 3 * cm])
    stats_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#0F4061")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements.append(stats_table)
    elements.append(Spacer(1, 12))

    # Gap details grouped by source
    gaps_by_source: dict[str, list] = {}
    for gap in result.get("gaps", []):
        src = gap.get("control_source", "")
        gaps_by_source.setdefault(src, []).append(gap)

    elements.append(Paragraph(
        "Gap Detalları" if lang == "az" else "Gap Details", heading_style
    ))
    for source, gaps in gaps_by_source.items():
        elements.append(Paragraph(_esc(source), sub_heading_style))

        compliant_n = sum(1 for g in gaps if g.get("status") == "compliant")
        partial_n = sum(1 for g in gaps if g.get("status") == "partial")
        missing_n = sum(1 for g in gaps if g.get("status") == "missing")
        summary_label = (
            f"{compliant_n} uyğun / {partial_n} qismən / {missing_n} yox"
            if lang == "az"
            else f"{compliant_n} compliant / {partial_n} partial / {missing_n} missing"
        )
        elements.append(Paragraph(summary_label, body_style))
        elements.append(Spacer(1, 4))

        for gap in gaps:
            status = gap.get("status", "missing")
            status_label = labels.get(status, status)
            color = STATUS_COLORS_PDF.get(status, colors.black)
            color_hex = f"#{color.hexval()[2:]}" if hasattr(color, "hexval") else "#000000"
            status_text = f'<font color="{color_hex}">[{_esc(status_label)}]</font>'
            elements.append(Paragraph(
                f"{status_text} <b>{_esc(gap.get('control_category', ''))}</b> "
                f"({_esc(gap.get('control_id', ''))})",
                body_style,
            ))

            if gap.get("control_text"):
                ctrl_label = "Nəzarət tələbi" if lang == "az" else "Control requirement"
                elements.append(Paragraph(
                    f"<b>{ctrl_label}:</b> {_esc(gap['control_text'])}", body_style
                ))

            if gap.get("standard_requirement"):
                req_label = "Standart tələbi" if lang == "az" else "Standard requirement"
                elements.append(Paragraph(
                    f"<b>{req_label}:</b> {_esc(gap['standard_requirement'])}", body_style
                ))

            cur_text = gap.get("current_document_text", "")
            if cur_text and cur_text != "Text not found":
                cur_label = "Sənəddəki mətn" if lang == "az" else "Current document text"
                elements.append(Paragraph(
                    f'<b>{cur_label}:</b> "{_esc(cur_text)}"', body_style
                ))

            if gap.get("gap_analysis"):
                gap_label = "Gap analizi" if lang == "az" else "Gap analysis"
                elements.append(Paragraph(
                    f"<b>{gap_label}:</b> {_esc(gap['gap_analysis'])}", body_style
                ))

            if gap.get("remediation_proposal"):
                rem_label = "Remediation təklifi" if lang == "az" else "Remediation proposal"
                elements.append(Paragraph(
                    f"<b>{rem_label}:</b> {_esc(gap['remediation_proposal'])}",
                    remediation_style,
                ))

            if gap.get("potential_risks"):
                pr_label = "Potensial risklər" if lang == "az" else "Potential risks"
                elements.append(Paragraph(
                    f"<b>{pr_label}:</b> {_esc(gap['potential_risks'])}",
                    body_style
                ))

            if gap.get("justification"):
                just_label = "Əsaslandırma" if lang == "az" else "Justification"
                elements.append(Paragraph(
                    f"<b>{just_label}:</b> {_esc(gap.get('justification', ''))}", body_style
                ))

            if gap.get("evidence_snippet"):
                ev_label = "Dəlil" if lang == "az" else "Evidence"
                ev_ref = gap.get("evidence_reference", "")
                ev_html = f"<i>{_esc(gap['evidence_snippet'])}</i>"
                if ev_ref:
                    ev_html += f" — {_esc(ev_ref)}"
                elements.append(Paragraph(ev_html, evidence_style))

            elements.append(Spacer(1, 6))

    # Recommendations
    elements.append(Paragraph(
        "Tövsiyələr" if lang == "az" else "Recommendations", heading_style
    ))
    for i, rec in enumerate(result.get("recommendations", []), 1):
        elements.append(Paragraph(f"{i}. {_esc(rec)}", body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
